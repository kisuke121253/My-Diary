from django.shortcuts import render, redirect
from django.contrib.auth.forms import AuthenticationForm
from django.contrib import messages
from django.views.decorators.http import require_POST
from django.shortcuts import redirect, get_object_or_404
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
import whisper
from django.core.files.base import ContentFile
from django.http import HttpResponse
import os
import re
from datetime import datetime
from babel.dates import format_date
from django.db.models.functions import TruncDate
from collections import OrderedDict
from itertools import groupby
import itertools
import google.generativeai as genai
import textwrap
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import Diario, Pagina
from django.contrib.auth import authenticate, login, logout
from .forms import SignUpForm, LoginForm
from django.contrib.auth.models import User
from django.utils import timezone
import pytz
from django.utils.timezone import localtime, activate
from django.utils.timezone import now
from .models import MoodLog
from .models import Evento, Nota
from .models import PageCount
import datetime
from django.shortcuts import redirect
from django.utils.dateparse import parse_datetime
from django.core.serializers import serialize
from django.utils.timezone import localtime
import json
from django.views.decorators.http import require_http_methods
from .models import Anotacao
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from datetime import date, timedelta
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER
from .models import AcompanhamentoSemanal, MoodLog
from django.db.models import Q
from .models import AcompanhamentoSemanal
from datetime import date
from datetime import date, timedelta
from django.db.models import Count


def to_markdown(text):
    text = text.replace('•', '  *')
    return textwrap.indent(text, '> ', predicate=lambda _: True)

def home(request):
    if request.user.is_authenticated:
        return redirect('diario')  
    else:
        return redirect('login') 

def registro(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data.get('username')
            raw_password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=raw_password)
            login(request, user)
            return redirect('diario') 
    else:
        form = SignUpForm()
    return render(request, 'registro.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('diario') 
        else:
            return render(request, 'login.html', {'form': form, 'error': 'Usuário ou senha inválidos'})
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})

def logout_view(request):
    logout(request)
    return redirect('login')  

@login_required
def calendario(request):
    activate(pytz.timezone('America/Sao_Paulo'))

    eventos = Evento.objects.filter(user=request.user, data_inicio__gte=now()).order_by('data_inicio')

    eventos_ajustados = [
        {
            'id': evento.id,
            'titulo': evento.titulo,
            'data_inicio': localtime(evento.data_inicio).strftime('%Hh%M'),  
            'data_fim': localtime(evento.data_fim).strftime('%Hh%M'),       
        } for evento in eventos
    ]

    proximos_eventos = eventos_ajustados[:5]

    return render(request, 'calendario.html', {
        'eventos': eventos_ajustados,
        'proximos_eventos': proximos_eventos
    })

@login_required
def detalhe_evento(request, evento_id):
    evento = get_object_or_404(Evento, pk=evento_id)
    return render(request, 'calendario.html', {'evento': evento})

@login_required
def diario(request):
    date_today = timezone.now().date()
    diario = Diario.objects.filter(user=request.user).first() 
    already_submitted = MoodLog.objects.filter(user=request.user, date=date_today).exists()
    context = {
        'diary_name': diario.nome if diario else None,
        'diarios': Diario.objects.filter(user=request.user).order_by('-data_criacao'),
        'already_submitted': already_submitted,
        'diario': diario  
    }

    return render(request, 'diario.html', context)

@login_required
def salvar_nome_diario(request):
    if request.method == "POST":
        nome_diario = request.POST.get('diary_name')
        diario, created = Diario.objects.get_or_create(user=request.user)
        diario.nome = nome_diario
        diario.save()

        return redirect('diario')
    else:
        
        return redirect('diario')
    

@login_required
@csrf_exempt 
def salvar_pagina(request):
    if request.method == 'POST':
        numero_pagina = request.POST.get('numero_pagina')
        conteudo = request.POST.get('conteudo')
        diario, _ = Diario.objects.get_or_create(user=request.user)  

        pagina, created = Pagina.objects.update_or_create(
            diario=diario,
            numero_pagina=numero_pagina,
            defaults={'conteudo': conteudo},
        )

        return JsonResponse({'status': 'sucesso'})
    else:
        return JsonResponse({'status': 'falha'}, status=400)
    
@login_required
def buscar_todas_paginas(request):
    try:
        diario = Diario.objects.get(user=request.user)
        paginas = diario.paginas.all().order_by('numero_pagina')
        dados_paginas = [{'numero_pagina': p.numero_pagina, 'conteudo': p.conteudo} for p in paginas]
        return JsonResponse({'status': 'sucesso', 'paginas': dados_paginas})
    except Diario.DoesNotExist:
        return JsonResponse({'status': 'erro', 'mensagem': 'Diário não encontrado'}, status=404)

@login_required
def editar_diario(request, diario_id):
    diario = get_object_or_404(Diario, id=diario_id, user=request.user)  
    if request.method == 'POST':
        diario.nome = request.POST.get('nome_diario')
        diario.save()
        return redirect('diario')  
    return render(request, 'editar_diario.html', {'diario': diario})

@login_required
def save_page_count(request):
    if request.method == 'POST':
        page_count = int(request.POST.get('page_count'))
        today = datetime.date.today()
        user = request.user

        existing_count = PageCount.objects.filter(user=user, date=today).first()
        if existing_count:
            existing_count.count = page_count
            existing_count.is_saved = True
            existing_count.save()
        else:
            PageCount.objects.create(user=user, count=page_count, date=today, is_saved=True)
        
        return JsonResponse({'status': 'success'})
    else:
        return JsonResponse({'status': 'error', 'message': 'Invalid request'}, status=400)

@login_required
def get_today_page_count(request):
    today_count = PageCount.objects.filter(user=request.user, date=datetime.date.today()).first()
    if today_count:
        return JsonResponse({'status': 'success', 'count': today_count.count})
    else:
        return JsonResponse({'status': 'error', 'message': 'No data found'}, status=404)

@require_POST
@login_required
def create_event(request):
    titulo = request.POST.get('titulo')
    descricao = request.POST.get('descricao')
    data_inicio = request.POST.get('data_inicio')
    data_fim = request.POST.get('data_fim')
    
    sao_paulo = pytz.timezone('America/Sao_Paulo')

    data_inicio = parse_datetime(data_inicio)
    if data_inicio:
        data_inicio = sao_paulo.localize(data_inicio)

    data_fim = parse_datetime(data_fim)
    if data_fim:
        data_fim = sao_paulo.localize(data_fim)
    
    Evento.objects.create(
        user=request.user,
        titulo=titulo,
        descricao=descricao,
        data_inicio=data_inicio,
        data_fim=data_fim
    )
    return redirect('calendario')


@login_required
def get_events(request):
    activate(pytz.timezone('America/Sao_Paulo'))

    eventos = Evento.objects.filter(user=request.user)

    eventos_data = []
    for evento in eventos:
        evento_dict = {
            'id': evento.id, 
            'title': evento.titulo,
            'start': localtime(evento.data_inicio).isoformat(),
            'end': localtime(evento.data_fim).isoformat(),
            'notas': list(evento.notas.all().values('id', 'texto', 'criado_em'))
        }
        eventos_data.append(evento_dict)

    return JsonResponse({'eventos': eventos_data})

api_key = 'AIzaSyASEelCOGxux50q-JQIIEoc9fWhxl-0mLE'
genai.configure(api_key=api_key)

@csrf_exempt
@require_http_methods(["POST"])
def get_recommendations(request):
    try:
        if not request.body:
            return JsonResponse({'error': 'No data provided'}, status=400)

        data = json.loads(request.body)
        
        user_prompt = data.get('prompt')
        if not user_prompt:
            return JsonResponse({'error': 'Prompt not provided'}, status=400)

        full_prompt = (
            "Você é um ajudante e deve ser curto e direto, recomende filmes, séries e livros e como eles podem ajudar com o problema"
            f"para ajudarem a pessoa com o seguinte problema: {user_prompt}"
        )

        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(full_prompt)
        answer = response.text

        return JsonResponse({'message': answer}, status=200)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
@require_http_methods(["POST"])
@login_required
def salvar_anotacao(request):
    try:
        data = json.loads(request.body)
        texto = data['texto']
        
        Anotacao.objects.create(usuario=request.user, texto=texto)
        
        return JsonResponse({'message': 'Anotação salva com sucesso!'}, status=200)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def buscar_anotacoes(request):
    anotacoes = Anotacao.objects.filter(usuario=request.user).order_by('-data_criacao')
    data = [{'id': anotacao.id, 'texto': anotacao.texto, 'data_criacao': anotacao.data_criacao.strftime('%Y-%m-%d %H:%M')} for anotacao in anotacoes]
    return JsonResponse(data, safe=False)  

@login_required
@require_http_methods(["DELETE"])
def excluir_anotacao(request, anotacao_id):
    try:
        anotacao = Anotacao.objects.get(id=anotacao_id, usuario=request.user)
        anotacao.delete()
        return JsonResponse({'message': 'Anotação excluída com sucesso!'}, status=200)
    except Anotacao.DoesNotExist:
        return JsonResponse({'error': 'Anotação não encontrada.'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def buscar_por_conteudo(request):
    query = request.GET.get('query')
    try:
        diario = Diario.objects.get(user=request.user)
        pagina = diario.paginas.filter(conteudo__icontains=query).first()
        if pagina:
            return JsonResponse({'status': 'sucesso', 'numero_pagina': pagina.numero_pagina})
        else:
            return JsonResponse({'status': 'falha', 'mensagem': 'Texto não encontrado'}, status=404)
    except Diario.DoesNotExist:
        return JsonResponse({'status': 'erro', 'mensagem': 'Diário não encontrado'}, status=404)


@login_required
def salvar_humor(request):
    if request.method == 'POST' and request.user.is_authenticated:
        mood = request.POST.get('mood')
        date = timezone.now().date()
        MoodLog.objects.update_or_create(user=request.user, date=date, defaults={'mood': mood})
        return redirect('diario')
    return redirect('login')  

@login_required
def acompanhamento_semanal_view(request):
    nome = request.user.username
    data = date.today().isoformat()
    numero_sessao = calcular_numero_sessao(request.user)

    if request.method == 'POST':
        obj = AcompanhamentoSemanal(
            nome=nome,
            data=data,
            numero_sessao=numero_sessao,
            avaliacao_humor=request.POST.get('avaliacao_humor', ''),
            mudancas_humor=request.POST.get('mudancas_humor', ''),
            eventos_semana=request.POST.get('eventos_semana', ''),
            sintomas=request.POST.get('sintomas', ''),
            estrategias_coping=request.POST.get('estrategias_coping', ''),
            observacoes_terapeuta=request.POST.get('observacoes_terapeuta', ''),
            objetivos_proxima_semana=request.POST.get('objetivos_proxima_semana', '')
        )
        obj.save()
        return redirect('acompanhamento')
    else:
        print('erro')

    one_week_ago = date.today() - timedelta(days=7)
    moods = MoodLog.objects.filter(
        user=request.user, 
        date__gte=one_week_ago
    ).values('mood').annotate(count=Count('mood'))

    mood_data = {mood['mood']: mood['count'] for mood in moods}

    all_moods = ['feliz', 'triste', 'irritado', 'ansioso', 'entusiasmado', 'cansado', 'estressado', 'indiferente', 'esperancoso', 'inspirado']
    for mood in all_moods:
        if mood not in mood_data:
            mood_data[mood] = 0

    mood_data_json = json.dumps(mood_data)

    return render(request, 'relatorio.html', {'nome': nome, 'data': data, 'numero_sessao': numero_sessao, 'mood_data': mood_data_json})

def calcular_numero_sessao(user):
    ultima_sessao = AcompanhamentoSemanal.objects.filter(nome=user.username).order_by('-numero_sessao').first()
    if ultima_sessao:
        return ultima_sessao.numero_sessao + 1
    return 1


@csrf_exempt
def analyze_sentiment(request):
    if request.method == 'POST':
        one_week_ago = date.today() - timedelta(days=7)
        moods = MoodLog.objects.filter(
            user=request.user, 
            date__gte=one_week_ago
        ).values('mood').annotate(count=Count('mood'))
        
        mood_data = {mood['mood']: mood['count'] for mood in moods}
        
        user_input = ", ".join([f"{mood}: {count}" for mood, count in mood_data.items()])
        
        result = process_sentiment_analysis(user_input)
        return JsonResponse({'result': result})
    return JsonResponse({'error': 'Request must be POST.'}, status=400)

def process_sentiment_analysis(user_input):
    full_prompt = (
        "Você é um ajudante e deve ser ***CURTO E DIRETO e no formato ABNT***, faça a analise geral e profunda dos humores baseado nos seguintes moods com observações: "
        f" esses são os moods da pessoa: {user_input}"
    )

    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(full_prompt)
    return response.text

@csrf_exempt
def analyze_mood_change(request):
    if request.method == 'POST':
        one_week_ago = date.today() - timedelta(days=7)
        moods = MoodLog.objects.filter(
            user=request.user,
            date__gte=one_week_ago
        ).order_by('date').values('date', 'mood')

        if not moods:
            return JsonResponse({'error': 'No mood data available for analysis.'}, status=400)

        user_input = ", ".join([f"{mood['date']}: {mood['mood']}" for mood in moods])
        mood_change_description = analyze_mood_changes_with_ai(user_input)
        
        return JsonResponse({'mudancas_humor': mood_change_description})
    else:
        return JsonResponse({'error': 'Request must be POST.'}, status=400)

def analyze_mood_changes_with_ai(user_input):
    full_prompt = (
        "Você é um ajudante e deve ser ***CURTO E DIRETO e no formato ABNT***, faça a analise geral e profunda sobre a mudança de humor no decorrer da semana nos seguintes moods com observações: "
        f" esses são os moods da pessoa: {user_input}"
    )

    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(full_prompt)
    return response.text

@login_required
def download_pdf_view(request, acompanhamento_id):
    acompanhamento = get_object_or_404(AcompanhamentoSemanal, id=acompanhamento_id, nome=request.user.username)
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="acompanhamento_semanal_{acompanhamento_id}.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    title_style = styles['Title']
    title_style.alignment = TA_CENTER

    content_style = styles['BodyText']
    content_style.spaceAfter = 12

    story = []

    story.append(Paragraph("Acompanhamento Semanal", title_style))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(f"Nome: {acompanhamento.nome}", content_style))
    story.append(Paragraph(f"Data: {acompanhamento.data.strftime('%d/%m/%Y')}", content_style))
    story.append(Paragraph(f"Número da Sessão: {acompanhamento.numero_sessao}", content_style))

    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Avaliação de Humor:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.avaliacao_humor, content_style))

    story.append(Paragraph("Mudanças de Humor:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.mudancas_humor, content_style))

    story.append(Paragraph("Eventos da Semana:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.eventos_semana, content_style))

    story.append(Paragraph("Sintomas:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.sintomas, content_style))

    story.append(Paragraph("Estratégias de Coping:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.estrategias_coping, content_style))

    story.append(Paragraph("Objetivos para a Próxima Semana:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.objetivos_proxima_semana, content_style))

    story.append(Paragraph("Observações do Terapeuta:", styles['Heading2']))
    story.append(Paragraph(acompanhamento.observacoes_terapeuta, content_style))

    doc.build(story)
    return response

@login_required
def download_pdf_view_criar(request):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="acompanhamento_semanal.pdf"'

    doc = SimpleDocTemplate(response, pagesize=letter,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    title_style = styles['Title']
    title_style.alignment = TA_CENTER

    content_style = styles['BodyText']
    content_style.spaceAfter = 12

    story = []

    story.append(Paragraph("Acompanhamento Semanal", title_style))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(f"Nome: {request.POST.get('nome', '')}", content_style))
    story.append(Paragraph(f"Data: {request.POST.get('data', '')}", content_style))
    story.append(Paragraph(f"Número da Sessão: {request.POST.get('numero_sessao', '')}", content_style))

    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Avaliação de Humor:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('avaliacao_humor', ''), content_style))

    story.append(Paragraph("Mudanças de Humor:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('mudancas_humor', ''), content_style))

    story.append(Paragraph("Eventos da Semana:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('eventos_semana', ''), content_style))

    story.append(Paragraph("Sintomas:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('sintomas', ''), content_style))

    story.append(Paragraph("Estratégias de Coping:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('estrategias_coping', ''), content_style))

    story.append(Paragraph("Objetivos para a Próxima Semana:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('objetivos_proxima_semana', ''), content_style))

    story.append(Paragraph("Observações do Terapeuta:", styles['Heading2']))
    story.append(Paragraph(request.POST.get('observacoes_terapeuta', ''), content_style))

    doc.build(story)
    return response

@login_required
def download_word_view(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="acompanhamento_semanal.docx"'

    doc = Document()

    title = doc.add_heading('Acompanhamento Semanal', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_section(doc, title_text, body_text):
        heading = doc.add_heading(title_text, level=2)
        paragraph = doc.add_paragraph(body_text)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(12)

    def clean_text(text):
        return text.replace('*', '')

    add_section(doc, 'Nome:', request.POST.get('nome', ''))
    add_section(doc, 'Data:', request.POST.get('data', ''))
    add_section(doc, 'Número da Sessão:', request.POST.get('numero_sessao', ''))
    add_section(doc, 'Avaliação de Humor:', clean_text(request.POST.get('avaliacao_humor', '')))
    add_section(doc, 'Mudanças de Humor:', clean_text(request.POST.get('mudancas_humor', '')))
    add_section(doc, 'Eventos da Semana:', clean_text(request.POST.get('eventos_semana', '')))
    add_section(doc, 'Sintomas:', request.POST.get('sintomas', ''))
    add_section(doc, 'Estratégias de Coping:', request.POST.get('estrategias_coping', ''))
    add_section(doc, 'Objetivos para a Próxima Semana:', request.POST.get('objetivos_proxima_semana', ''))
    add_section(doc, 'Observações do Terapeuta:', request.POST.get('observacoes_terapeuta', ''))

    doc.save(response)
    return response

@csrf_exempt
@login_required
def consultar_eventos(request):
    if request.method == 'POST':
        data = json.loads(request.body)
       
        hoje = datetime.date.today()
        inicio_semana = hoje - datetime.timedelta(days=hoje.weekday())
        fim_semana = inicio_semana + datetime.timedelta(days=6)
        
        eventos = Evento.objects.filter(user=request.user, data_inicio__gte=inicio_semana, data_fim__lte=fim_semana)
        eventos_formatados = [evento.titulo for evento in eventos]
     
        full_prompt = (
            "Você é um ajudante e deve ser ***CURTO E DIRETO e no formato ABNT***, faça a análise geral dos profunda sobre os eventos ocorridos no decorrer da semana da pessoa: "
            f" esses são os eventos da pessoa: {eventos_formatados}"
        )
        
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(full_prompt)
        analise_ia = response.text 
        

        return JsonResponse({
            'success': True,
            'eventos': eventos_formatados,
            'analise_ia': analise_ia
        })

    return JsonResponse({'success': False, 'error': 'Invalid request'})


@login_required
def list_acompanhamento_semanal(request):
    acompanhamentos = AcompanhamentoSemanal.objects.filter(nome=request.user.username)

    filter_date = request.GET.get('filter_date')
    if filter_date:
        acompanhamentos = acompanhamentos.filter(data=filter_date)

    last_report = acompanhamentos.order_by('-data').first()

    can_generate_report = True
    days_until_next_report = 0  

    if last_report and (timezone.now().date() - last_report.data).days < 7:
        can_generate_report = False
        days_until_next_report = 7 - (timezone.now().date() - last_report.data).days

    return render(request, 'list_acompanhamento.html', {
        'acompanhamentos': acompanhamentos,
        'can_generate_report': can_generate_report,
        'days_until_next_report': days_until_next_report
    })

@login_required
def delete_acompanhamento(request, id):
    try:
        acompanhamento = AcompanhamentoSemanal.objects.get(pk=id)
        acompanhamento.delete()
        messages.success(request, 'Acompanhamento deletado com sucesso!')
    except AcompanhamentoSemanal.DoesNotExist:
        messages.error(request, 'Acompanhamento não encontrado.')
    return redirect('acompanhamento') 

@csrf_exempt
@login_required
def add_note(request, event_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        evento = Evento.objects.get(id=event_id, user=request.user)
        nova_nota = Nota(evento=evento, texto=data['texto'])
        nova_nota.save()
        return JsonResponse({'status': 'success', 'note_id': nova_nota.id})
    return JsonResponse({'status': 'error'}, status=400)

@csrf_exempt
@require_POST
def delete_note(request, note_id):
    try:
        nota = Nota.objects.get(id=note_id)
        nota.delete()
        return JsonResponse({'status': 'success'}, status=200)
    except Nota.DoesNotExist:
        return JsonResponse({'status': 'not found'}, status=404)

@login_required
def view_acompanhamento(request, id):
    acompanhamento = get_object_or_404(AcompanhamentoSemanal, id=id, nome=request.user.username)
    data_referencia = acompanhamento.data 
    data_inicio = data_referencia - timedelta(days=7)  

    avaliacoes_humor = MoodLog.objects.filter(
        user=request.user,
        date__range=(data_inicio, data_referencia)
    ).values_list('mood', flat=True)

    humor_counts = {
        'feliz': 0,
        'triste': 0,
        'irritado': 0,
        'ansioso': 0,
        'entusiasmado': 0,
        'cansado': 0,
        'estressado': 0,
        'indiferente': 0,
        'esperancoso': 0,
        'inspirado': 0
    }

    for avaliacao in avaliacoes_humor:
        for humor in avaliacao.split(','):
            if humor in humor_counts:
                humor_counts[humor] += 1

    mood_data_json = json.dumps(humor_counts)

    if request.method == 'POST':
        acompanhamento.sintomas = request.POST.get('sintomas', '')
        acompanhamento.estrategias_coping = request.POST.get('estrategias_coping', '')
        acompanhamento.objetivos_proxima_semana = request.POST.get('objetivos_proxima_semana', '')
        acompanhamento.save()
        return JsonResponse({'success': True})

    return render(request, 'view_acompanhamento.html', {
        'acompanhamento': acompanhamento,
        'mood_data': mood_data_json
    })